from openai import OpenAI
from docx import Document
from typing import Dict
import glob
from pprint import pprint

client = OpenAI()

def transcribe_audio(audio_file_path: str) -> str:
    with open(audio_file_path, 'rb') as audio_file:
        try:
            transcription = client.audio.transcriptions.create(model='whisper-1', file=audio_file)
        except Exception as e:
            print(e)
            from pprint import pprint
            pprint(e.__dict__)
            raise
        
    return transcription.text

def abstract_summary_extraction(transcription: str) -> str:
    response = client.chat.completions.create(
        model='gpt-4', 
        temperature=0,
        messages=[
            {
                'role': 'system', 
                'content': 'You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.'
            }, 
            {
                'role': 'user', 
                'content': transcription
            }
        ]
    )
    
    return response.choices[0].message.content

def key_points_extraction(transcription: str) -> str:
    response = client.chat.completions.create(
        model='gpt-4', 
        temperature=0,
        messages=[
            {
                'role': 'system', 
                'content': 'You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.'
            }, 
            {
                'role': 'user', 
                'content': transcription
             }
        ]
    )
    
    return response.choices[0].message.content

def action_item_extraction(transcription: str) -> str:
    response = client.chat.completions.create(
        model='gpt-4', 
        temperature=0,
        messages=[
            {
                'role': 'system', 
                'content': 'You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.'
            }, 
            {
                'role': 'user', 
                'content': transcription
            }
        ]
    )
    
    return response.choices[0].message.content

def sentiment_analysis(transcription: str) -> str:
    response = client.chat.completions.create(
       model='gpt-4', 
       temperature=0,
       messages=[
           {
               'role': 'system', 
               'content': 'As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.'
           }, 
           {
               'role': 'user', 
               'content': transcription
           }
       ] 
    )
    
    return response.choices[0].message.content

def meeting_minutes(transcription: str) -> str:
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    
    return {
        'abstract_summary': abstract_summary, 
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def save_as_docx(minutes: Dict[str, str], filename: str) -> None:
    doc = Document()
    
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading)
        doc.add_paragraph(value)
        doc.add_paragraph()
        
    doc.save(filename)
    
transcriptions = [transcribe_audio(audio_file_path) for audio_file_path in sorted(glob.glob('out*'))]
transcription = ''.join(transcriptions)
minutes = meeting_minutes(transcription)
pprint(minutes)
save_as_docx(minutes, 'meeting_minutes.docx')
